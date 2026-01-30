# docx_exporter.py
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def is_heading(line: str) -> bool:
    """Определяем, является ли строка заголовком"""
    line = line.strip()
    
    # Проверяем паттерны заголовков
    patterns = [
        r'^#+\s+.+',  # Markdown заголовки
        r'^[A-ZА-Я][^.!?]{0,50}[.!?]$',  # Короткие предложения без строчных букв
        r'^(Глава|Раздел|Часть|Параграф|§)\s+',  # Структурные элементы
        r'^\d+[\.\)]\s+[A-ZА-Я]',  # Нумерованные заголовки
    ]
    
    for pattern in patterns:
        if re.match(pattern, line, re.IGNORECASE):
            return True
    
    # Проверяем длину и наличие заглавных букв
    if len(line) < 100 and line[0].isupper():
        words = line.split()
        if len(words) < 10:
            return True
    
    return False

def is_list_item(line: str) -> bool:
    """Определяем, является ли строка элементом списка"""
    line = line.strip()
    
    # Паттерны списков
    patterns = [
        r'^[•\-*]\s+',  # Маркированные списки
        r'^\d+[\.\)]\s+',  # Нумерованные списки
        r'^[a-zA-Zа-яА-Я][\.\)]\s+',  # Алфавитные списки
    ]
    
    for pattern in patterns:
        if re.match(pattern, line):
            return True
    
    return False

def is_bullet_list(line: str) -> bool:
    """Определяем, является ли список маркированным"""
    line = line.strip()
    return bool(re.match(r'^[•\-*]\s+', line))

def smart_docx_export(text: str, original_mode: str = None) -> Document:
    """
    Умное сохранение в DOCX:
    - Определяет заголовки
    - Сохраняет списки
    - Добавляет метаданные
    """
    doc = Document()
    
    # Настройка стилей
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Заголовок документа
    if original_mode:
        title = doc.add_heading('Обработанный текст', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Метаданные
        meta = doc.add_paragraph()
        meta.add_run(f"Режим обработки: {original_mode}")
        meta.add_run(f"\nДата создания: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Пустая строка
    
    # Анализируем текст
    lines = text.split('\n')
    current_list = None  # Текущий тип списка
    
    for i, line in enumerate(lines):
        line = line.strip()
        
        if not line:  # Пустая строка
            current_list = None
            doc.add_paragraph()
            continue
        
        # Проверяем, является ли это заголовком
        if is_heading(line):
            current_list = None
            
            # Очищаем от маркеров заголовков
            clean_line = re.sub(r'^#+\s+', '', line)
            clean_line = re.sub(r'^\d+[\.\)]\s+', '', clean_line)
            
            # Определяем уровень заголовка
            if re.match(r'^#\s+', line):
                level = 1
            elif re.match(r'^##\s+', line):
                level = 2
            elif re.match(r'^###\s+', line):
                level = 3
            else:
                # Автоматически определяем уровень по контексту
                if i == 0 or not lines[i-1].strip():
                    level = 1
                else:
                    level = 2
            
            doc.add_heading(clean_line, level=level)
        
        # Проверяем, является ли это списком
        elif is_list_item(line):
            list_type = 'bullet' if is_bullet_list(line) else 'number'
            
            # Если тип списка изменился
            if current_list != list_type:
                current_list = list_type
            
            # Очищаем от маркеров списка
            clean_line = re.sub(r'^[•\-*]\s+', '', line)
            clean_line = re.sub(r'^\d+[\.\)]\s+', '', clean_line)
            clean_line = re.sub(r'^[a-zA-Zа-яА-Я][\.\)]\s+', '', clean_line)
            
            if list_type == 'bullet':
                p = doc.add_paragraph(style='List Bullet')
            else:
                p = doc.add_paragraph(style='List Number')
            
            p.add_run(clean_line)
        
        # Обычный текст
        else:
            current_list = None
            
            # Проверяем, не является ли это продолжением предыдущего абзаца
            if i > 0 and lines[i-1].strip() and not is_heading(lines[i-1]) and not is_list_item(lines[i-1]):
                # Добавляем к последнему абзацу
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.add_run(' ' + line)
            else:
                # Новый абзац
                p = doc.add_paragraph(line)
    
    return doc

# Импортируем datetime для метаданных
from datetime import datetime